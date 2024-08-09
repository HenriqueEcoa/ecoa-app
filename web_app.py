from flask import Flask, render_template, request, redirect, url_for, jsonify, send_file,flash
from flask_login import LoginManager, UserMixin, login_user, login_required, logout_user, current_user
from werkzeug.security import generate_password_hash, check_password_hash
import sqlite3
import openai  
import json
import os
import io
import dotenv
import smtplib
import email.message
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement



app = Flask(__name__)

app.secret_key = 'supersecretkey'  # Use uma chave secreta mais segura em produção


login_manager = LoginManager()
login_manager.init_app(app)

class User(UserMixin):
    def __init__(self, id, username):
        self.id = id
        self.username = username



def get_db_connection():
    conn = sqlite3.connect('database.db')
    conn.row_factory = sqlite3.Row
    return conn


@login_manager.user_loader
def load_user(user_id):
    conn = get_db_connection()
    user_data = conn.execute('SELECT id, username FROM users WHERE id = ?', (user_id,)).fetchone()
    conn.close()
    if user_data:
        return User(user_data['id'], user_data['username'])
    return None


# Arquivo .env
dotenv.load_dotenv()

openai.api_key = os.getenv('OPENAI_KEY')

#Funçoes Complementares

#Funcão Chat GPT
def get_completion(prompt, model="gpt-3.5-turbo"):
    response = openai.ChatCompletion.create(
        model=model,
        messages=[{"role": "user", "content": prompt}]
    )
    return response.choices[0].message.content.strip()

# Add ideia ao banco de dados
def add_idea_to_db(user_id, idea=None, audio_file=None, transcription=None, titulo=None, descricao=None, impacto=None, recursos=None, beneficios=None, observacao=None, resumo=None, data_inicio=None, data_implementacao=None,visibilidade=None):
    print(f"Adicionando ideia: user_id={user_id}, idea={idea}, audio_file={audio_file}, transcription={transcription}, titulo={titulo}, descricao={descricao}, impacto={impacto}, recursos={recursos}, beneficios={beneficios}, observacao={observacao}, resumo={resumo}, data_inicio={data_inicio}, data_implementacao={data_implementacao},visibilidade={visibilidade}")
    try:
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        c.execute('''INSERT INTO ideas (user_id, idea, audio_file, transcription, titulo, descricao, impacto, recursos, beneficios, observacao, resumo, data_inicio, data_implementacao,visibilidade) 
                     VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)''', 
                  (user_id, idea, audio_file, transcription, titulo, descricao, impacto, recursos, beneficios, observacao, resumo, data_inicio, data_implementacao,visibilidade))
        conn.commit()
        print("Ideia adicionada com sucesso!")
        return True
    except sqlite3.Error as e:
        print(f"Erro ao adicionar ideia no banco de dados: {e}")
    finally:
        conn.close()
    

# Função para enviar email

def enviar_email():
    corpo = """
        <img style="width:840px" src="https://emaillogo.netlify.app/WhatsApp%20Image%202024-07-29%20at%2017.43.37.jpeg"/>
    """

    msg = email.message.Message()
    msg["Subject"] = "Lembrete do Capit"
    msg["From"] = "unoecoa@gmail.com" #Email Remetente
    msg["To"] = "jaccoudmarcela@gmail.com" #Email Destinatario

    # Senha do app email /google ->  /unoecoa@gmail.com/
    password = "fhlr esyf yjze mkwk"

    msg.add_header("Content-Type", "text/html")
    msg.set_payload(corpo)

    s = smtplib.SMTP("smtp.gmail.com: 587")
    s.starttls()
    s.login(msg["From"], password)
    s.sendmail(msg["From"], [msg["To"]], msg.as_string().encode("utf-8"))

    print("Email Enviado")



#APP
@app.route('/minha_ideias')
@login_required
def minha_ideias():
    with sqlite3.connect('database.db') as conn:
        c = conn.cursor()

        # Consultar todas as ideias do usuário
        c.execute("SELECT * FROM ideas WHERE user_id = ?", (current_user.username,))
        ideas = c.fetchall()  # Use fetchall() para obter todas as linhas

        # Consultar todas as tags
        c.execute('SELECT id, name FROM tags')
        tags = c.fetchall()
    
    # Verificar o conteúdo de `ideas` e `tags` para depuração
    print("Ideias do usuário:", ideas,current_user.username)
    print("Tags:", tags)

    return render_template('ideia.html', ideas=ideas, tags=tags,current_user=current_user)



@app.route('/update/<int:id>', methods=['POST'])
@login_required
def update_idea(id):
    try:
        form_data = request.form
        print("Dados do formulário recebidos:", form_data)

        user_id = form_data.get('user_id')
        titulo = form_data.get('titulo')
        descricao = form_data.get('descricao')
        impacto = form_data.get('impactos')
        recursos = form_data.get('recursos')
        beneficios = form_data.get('beneficios')
        resumo = form_data.get('resumo')
        data_inicio = form_data.get('data_inicio')
        data_implementacao = form_data.get('data_implementacao')
        observacao = form_data.get('observacoes')
        new_tags = form_data.get('tags').split(',')
        visibilidade = form_data.get('visibilidade')

        with sqlite3.connect('database.db') as conn:
            c = conn.cursor()
            
            c.execute("""
                UPDATE ideas
                SET observacao=?, user_id=?, titulo=?, descricao=?, impacto=?, recursos=?, beneficios=?, 
                    resumo=?, data_inicio=?, data_implementacao=?,visibilidade=?
                WHERE id=?
                """, (observacao, user_id, titulo, descricao, impacto, recursos, beneficios, 
                      resumo, data_inicio, data_implementacao,visibilidade, id))
            
            c.execute("DELETE FROM idea_tags WHERE idea_id = ?", (id,))

            for tag in new_tags:
                tag = tag.strip()
                if tag:
                    c.execute("SELECT id FROM tags WHERE name = ?", (tag,))
                    tag_id = c.fetchone()
                    if tag_id:
                        tag_id = tag_id[0]
                    else:
                        c.execute("INSERT INTO tags (name) VALUES (?)", (tag,))
                        tag_id = c.lastrowid
                        
                    c.execute("INSERT INTO idea_tags (idea_id, tag_id) VALUES (?, ?)", (id, tag_id))
            
            conn.commit()

        return redirect(url_for('estruturar', idea_id=id))
    
    except sqlite3.Error as e:
        return "Erro ao acessar o banco de dados", 500
    except Exception as e:
        print(f"Erro ao processar o formulário: {e}")
        return render_template('erro.html')




def get_idea_by_id(idea_id):
    try:
        with sqlite3.connect('database.db') as conn:
            c = conn.cursor()
            c.execute("SELECT * FROM ideas WHERE id=?", (idea_id,))
            idea = c.fetchone()
        return idea
    except sqlite3.Error as e:
        print(f"Erro ao obter ideia do banco de dados: {e}")
        return None

# Rota para exibir a ideia
@app.route('/ideia_pagina/<int:idea_id>')
@login_required
def view_idea(idea_id):
    color=request.args.get('color')
    
    with sqlite3.connect('database.db') as conn:
        c = conn.cursor()
        c.execute("SELECT * FROM ideas WHERE user_id = ? ", (current_user.username,))
        ideas = c.fetchall()

    idea = get_idea_by_id(idea_id)
    if idea:
        return render_template('pagina_ideia.html', idea=idea, ideas=ideas,color=color)
    else:
        return render_template('erro.html')


def get_idea_id(user_id, idea):
    try:
        with sqlite3.connect('database.db') as conn:
            c = conn.cursor()
            c.execute("""
                SELECT id FROM ideas 
                WHERE user_id=? AND idea=? 
                ORDER BY id DESC LIMIT 1
            """, (user_id, idea))
            result = c.fetchone()
            if result:
                return result[0]
            else:
                return None
    except sqlite3.Error as e:
        print(f"Erro ao obter ID da ideia do banco de dados: {e}")
        return None

def add_tag_to_db(tag_name):
    try:
        with sqlite3.connect('database.db') as conn:
            cursor = conn.cursor()
            cursor.execute('INSERT OR IGNORE INTO tags (name) VALUES (?)', (tag_name,))
            conn.commit()
            cursor.execute('SELECT id FROM tags WHERE name = ?', (tag_name,))
            tag_id = cursor.fetchone()[0]
            return tag_id
    except sqlite3.Error as e:
        print(f"Erro ao adicionar a tag ao banco de dados: {e}")
        return None

def associate_tag_with_idea(idea_id, tag_id):
    try:
        with sqlite3.connect('database.db') as conn:
            cursor = conn.cursor()
            cursor.execute('INSERT INTO idea_tags (idea_id, tag_id) VALUES (?, ?)', (idea_id, tag_id))
            conn.commit()
    except sqlite3.Error as e:
        print(f"Erro ao associar a tag com a ideia: {e}")


@app.route('/add', methods=['POST'])
@login_required
def add():
    if request.method == 'POST':
        user_id = request.form['user_id']
        idea = request.form['idea']
        tags = request.form.get('tags', '[]') 
        status = '0'
        visibilidade = request.form.get('visibilidade')
        

        try:
            tags_list = json.loads(tags)
        except json.JSONDecodeError:
            tags_list = []

        # Enviar email
        enviar_email()
        

        prompt = f"""
        Ajude na construção e fomentação da ideia: "{idea}", além disso ela deve responder os seguintes questionamentos e retornar no formato JSON:
        {{
            "descricao": "Por favor, descreva a ideia em detalhes. Certifique-se de que a descrição tenha no mínimo 100 palavras.",
            "impacto": "Se essa ideia for aplicada, qual será o impacto gerado? Detalhe os possíveis resultados e efeitos dessa implementação.",
            "recursos": "Quais são os recursos necessários para implementar essa ideia? Responda de forma detalhada, sem usar tópicos ou listas.",
            "beneficios": "Quais são os benefícios que essa ideia pode trazer? Descreva-os de forma contínua e sem listar em tópicos.",
            "resumo": "Com base nas informações fornecidas nas seções anteriores, incluindo a descrição da ideia e as respostas detalhadas, crie um resumo completo de pelo menos 300 palavras.",
            "titulo": "Com base em todas as informações fornecidas, sugira um título adequado que resuma a ideia e seu impacto com no máximo 6 palavras."
        }}
        """
        response = get_completion(prompt)

        try:
            response_json = json.loads(response)
            titulo = response_json.get("titulo")
            descricao = response_json.get("descricao")
            impacto = response_json.get("impacto")
            recursos = response_json.get("recursos")
            beneficios = response_json.get("beneficios")
            resumo = response_json.get("resumo")

        except json.JSONDecodeError:
            return render_template('erro.html')
        
        
        if add_idea_to_db(user_id, idea, None, None, titulo, descricao, impacto, recursos, beneficios, None, resumo, None, None,visibilidade): 
            idea_id = get_idea_id(user_id, idea)


        

            if idea_id is not None:
                for tag in tags_list:
                    tag_value = tag.get('value', '')
                    if tag_value:
                        tag_id = add_tag_to_db(tag_value)
                        if tag_id:
                            associate_tag_with_idea(idea_id, tag_id)
            
                with sqlite3.connect('database.db') as conn:
                    cursor = conn.cursor()
                    cursor.execute("SELECT * FROM ideas WHERE user_id = ? ", (current_user.username,))
                    ideas = cursor.fetchall()
                    cursor.execute('SELECT id, name FROM tags')
                    tags = cursor.fetchall()

                return render_template('ideia.html', ideas=ideas, tags=tags,status=status)

            else:
                return render_template('erro.html')
        else:
            return render_template('erro.html')



def search_ideas(keyword):
    try:
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        
        c.execute("""
            SELECT * FROM ideas
            WHERE titulo LIKE ?
            AND visibilidade = 'Público'
        """, ('%' + keyword + '%',))
        
        ideas = c.fetchall()
        
        conn.close()
        return ideas
    except sqlite3.Error as e:
        print(f"Erro ao buscar ideias: {e}")
        return []


@app.route('/search', methods=['GET'])
@login_required
def search():
    keyword = request.args.get('keyword') 
    if keyword:
        ideas = search_ideas(keyword)
        if ideas:
            return render_template('pesquisa.html',ideas=ideas,keyword=keyword)
        else:
            return render_template('pesquisa_erro.html')
    else:
        return render_template('erro.html')


def search_minhas_ideas(keyword):
    try:
        conn = sqlite3.connect('database.db')
        c = conn.cursor()
        
        c.execute("""
            SELECT * FROM ideas
            WHERE titulo LIKE ?
            AND user_id=?
        """, ('%' + keyword + '%',current_user.username))
        
        ideas = c.fetchall()
        
        conn.close()
        return ideas
    except sqlite3.Error as e:
        print(f"Erro ao buscar ideias: {e}")
        return []


@app.route('/search_minha_ideia', methods=['GET'])
@login_required
def search_minha_ideia():
    keyword = request.args.get('keyword') 
    if keyword:
        ideas = search_minhas_ideas(keyword)
        if ideas:
            return render_template('pesquisa.html',ideas=ideas,keyword=keyword)
        else:
            return render_template('pesquisa_erro.html')
    else:
        return render_template('erro.html')


#Como posso bloquear para n deixar as pessoas entrarem em URLS que n sao delas (Pergunta pro Bruno Curiosidade)
@app.route('/estruturar/<int:idea_id>')
@login_required
def estruturar(idea_id):
    with sqlite3.connect('database.db') as conn:
        c = conn.cursor()

        c.execute("SELECT * FROM users")
        users = c.fetchall()

        c.execute("SELECT * FROM ideas WHERE user_id = ? ", (current_user.username,))
        ideas = c.fetchall()

        c.execute(""" SELECT t.id, t.name FROM tags t INNER JOIN idea_tags it 
        ON t.id = it.tag_id WHERE it.idea_id = ? """, (idea_id,))
        tags = c.fetchall()

    idea = get_idea_by_id(idea_id)
    if idea:
        return render_template('estruturar.html', idea=idea,tags=tags,users=users)
    else:
        return render_template('erro.html')


def get_document(doc_id):
    document_data = {'titulo': '','resumo': '', 'descricao': '', 'impacto': '', 'recursos': '', 'beneficios': '', 'observacao': ''}
    try:
        with sqlite3.connect('database.db') as conn:
            c = conn.cursor()
            # Recuperar todos os campos do banco de dados
            c.execute("SELECT titulo, resumo, descricao, impacto, recursos, beneficios, observacao FROM ideas WHERE id=?", (doc_id,))
            result = c.fetchone()
            if result:
                # Preencher o dicionário com os valores recuperados
                document_data['titulo'], document_data['resumo'], document_data['descricao'], document_data['impacto'], document_data['recursos'], document_data['beneficios'], document_data['observacao'] = result
    except sqlite3.Error as e:
        print(f"Erro ao obter ideias do banco de dados: {e}")
    return document_data


def formatar_documento_abnt(doc_id):
    document_data = get_document(doc_id)

    if not any(document_data.values()):
        return None
    
    doc = Document()

    estilo_paragrafo = doc.styles['Normal']
    estilo_paragrafo.font.name = 'Times New Roman'
    estilo_paragrafo.font.size = Pt(12)
    estilo_paragrafo.paragraph_format.space_before = Pt(6)
    estilo_paragrafo.paragraph_format.space_after = Pt(6)
    estilo_paragrafo.paragraph_format.line_spacing = 1.5
    estilo_paragrafo.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    for key, value in document_data.items():
        if value:
            doc.add_paragraph(f"{key.capitalize()}:")
            doc.add_paragraph(value, style='Normal')
            doc.add_paragraph()  

    secao = doc.sections[0]
    secao.left_margin = Pt(72)
    secao.right_margin = Pt(72)
    secao.top_margin = Pt(72)
    secao.bottom_margin = Pt(72)

    static_dir = os.path.join(os.getcwd(), 'static')
    if not os.path.exists(static_dir):
        os.makedirs(static_dir)

    nome_arquivo = os.path.join(static_dir, f"documento_{doc_id}.docx")
    doc.save(nome_arquivo)

    return nome_arquivo

@app.route('/baixar', methods=['GET', 'POST'])
@login_required
def baixar():
    if request.method == 'POST':
        doc_id = request.form['id']
        documento_salvo = formatar_documento_abnt(doc_id)
        if documento_salvo and os.path.exists(documento_salvo):
            return send_file(documento_salvo, as_attachment=True)
    return render_template('erro.html')


def delete_tags(idea_id, conn):
    c = conn.cursor()
    c.execute("SELECT tag_id FROM idea_tags WHERE idea_id = ?", (idea_id,))
    tag_ids = c.fetchall()
    c.execute("DELETE FROM idea_tags WHERE idea_id = ?", (idea_id,))

    for tag_id in tag_ids:
        tag_id = tag_id[0]
        c.execute("SELECT COUNT(*) FROM idea_tags WHERE tag_id = ?", (tag_id,))
        count = c.fetchone()[0]
        
        if count == 0:
            c.execute("DELETE FROM tags WHERE id = ?", (tag_id,))

    conn.commit()

@app.route('/delete/<int:id>', methods=['POST'])
@login_required
def delete_idea(id):
    try:
        with sqlite3.connect('database.db') as conn:
            c = conn.cursor()
            delete_tags(id, conn)
            c.execute("DELETE FROM ideas WHERE id=?", (id,))
            conn.commit()
            ideas = c.fetchall()

        return redirect(url_for('minha_ideias', ideas=ideas))
    
    except Exception as e:
        print(f"Erro ao deletar a ideia: {e}")
        return str(e), 400


def get_tag_ids_by_names(tag_names):
    with sqlite3.connect('database.db') as conn:
        cursor = conn.cursor()
        query = '''
            SELECT id, name FROM tags WHERE name IN ({})
        '''.format(','.join('?' for _ in tag_names))
        cursor.execute(query, tag_names)
        rows = cursor.fetchall()
    tag_ids = [row[0] for row in rows]
    found_tag_names = [row[1] for row in rows]
    print(f"Tags encontradas no banco de dados: {found_tag_names}")
    return tag_ids


def search_ideas_by_tags(tag_ids):
    query = '''
        SELECT DISTINCT ideas.id, ideas.titulo, ideas.descricao
        FROM ideas
        JOIN idea_tags ON ideas.id = idea_tags.idea_id
        WHERE idea_tags.tag_id IN ({})
    '''.format(','.join('?' for _ in tag_ids))

    with sqlite3.connect('database.db') as conn:
        cursor = conn.cursor()
        cursor.execute(query, tag_ids)
        return cursor.fetchall()
    
@app.route('/search_tags', methods=['POST'])
@login_required
def search_by_tags():
    tag_names = request.form.getlist('tags')
    print(f"Tags recebidas: {tag_names}")
    
    if tag_names:
        tag_ids = get_tag_ids_by_names(tag_names)
        print(f"IDs das tags encontradas: {tag_ids}")
        if tag_ids:
            ideas = search_ideas_by_tags(tag_ids)
            if ideas:
                return render_template('pesquisa2.html', ideas=ideas,tag_names=tag_names)
            else:
                return render_template('pesquisa_erro.html')
        else:
            return render_template('pesquisa_erro.html')
    else:
        return render_template('erro.html')



@app.route('/erro')
def erro():
    return render_template('erro.html')


@app.route('/pesquisa_erro')
@login_required
def pesquisa_erro():
    return render_template('pesquisa_erro.html')



@app.route('/', methods=['GET', 'POST'])
def index2():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']

        conn = get_db_connection()
        user = conn.execute('SELECT id, username, password FROM users WHERE username = ?', (username,)).fetchone()
        conn.close()

        if user and check_password_hash(user['password'], password):
            user_obj = User(user['id'], user['username'])
            login_user(user_obj)
            return redirect(url_for('ideia_colab'))
        else:
            return 'Invalid credentials'

    return render_template('login.html')


##AMBIENTE COLABORATIVO##

@app.route('/ideia')
@login_required
def ideia_colab():
    with sqlite3.connect('database.db') as conn:
        c = conn.cursor()
        c.execute("SELECT * FROM ideas WHERE visibilidade = 'Público'")
        ideas = c.fetchall()

        c.execute('SELECT id, name FROM tags')
        tags = c.fetchall()
    return render_template('ideia_colab.html',ideas=ideas,tags=tags)

#Pagina de ideia colaborativo
@app.route('/ideia_pagina_colab/<int:idea_id>')
@login_required
def view_idea_colab(idea_id):
    color=request.args.get('color')
    
    with sqlite3.connect('database.db') as conn:
        c = conn.cursor()
        c.execute("SELECT * FROM ideas WHERE visibilidade = 'Público'")
        ideas = c.fetchall()

    idea = get_idea_by_id(idea_id)
    if idea:
        return render_template('pagina_ideia_colab.html', idea=idea, ideas=ideas, color=color)
    else:
        return render_template('erro.html')
    




@app.route('/perfil')
@login_required
def perfil():
    conn = get_db_connection()
    user = conn.execute('SELECT * FROM users WHERE id = ?', (current_user.id,)).fetchone()
    conn.close()
    return render_template('perfil.html', user=user)


@app.route('/update_profile', methods=['POST'])
@login_required
def update_profile():
    nome = request.form.get('nome', '').strip()
    sobrenome = request.form.get('sobrenome', '').strip()
    cargo = request.form.get('cargo', '').strip()
    setor = request.form.get('setor', '').strip()
    senha = request.form.get('senha', '').strip()
    confirmar_senha = request.form.get('confirmar_senha', '').strip()
    email = request.form.get('email', '').strip()

    if senha and senha != confirmar_senha:
        return jsonify({'success': False, 'error': 'A confirmação da senha não corresponde.'})

    conn = get_db_connection()
    
    update_fields = []
    params = []

    if nome:
        update_fields.append('nome = ?')
        params.append(nome)
    if sobrenome:
        update_fields.append('sobrenome = ?')
        params.append(sobrenome)
    if cargo:
        update_fields.append('cargo = ?')
        params.append(cargo)
    if setor:
        update_fields.append('setor = ?')
        params.append(setor)
    if email:
        update_fields.append('email = ?')
        params.append(email)

    if senha:
        hashed_password = generate_password_hash(senha, method='pbkdf2:sha256')
        update_fields.append('password = ?')
        params.append(hashed_password)

    if update_fields:
        query = 'UPDATE users SET ' + ', '.join(update_fields) + ' WHERE id = ?'
        params.append(current_user.id)
        conn.execute(query, params)
        conn.commit()

    conn.close()
    return redirect(url_for('perfil'))

@app.route('/upload_profile_picture', methods=['POST'])
@login_required
def upload_profile_picture():
    if 'profile_picture' in request.files:
        file = request.files['profile_picture']
        if file and file.mimetype.startswith('image/'):
            file_blob = file.read()
            conn = get_db_connection()
            conn.execute('UPDATE users SET profile_picture = ? WHERE id = ?', (file_blob, current_user.id))
            conn.commit()
            conn.close()
            return redirect(url_for('perfil'))
    flash('Erro ao carregar a imagem. Tente novamente.')
    return redirect(url_for('perfil'))


@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        name = request.form['name']
        setor = request.form['setor']

        hashed_password = generate_password_hash(password, method='pbkdf2:sha256')

        conn = get_db_connection()
        try:
            conn.execute('INSERT INTO users (username, password, nome, setor) VALUES (?, ?, ?, ?)', (username, hashed_password, name, setor))
            conn.commit()
            return redirect(url_for('index2'))
        except sqlite3.IntegrityError:
            return 'Username already exists'
        finally:
            conn.close()

    return render_template('registro.html')


@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('index2'))


@app.route('/get_profile_picture/<int:user_id>')
@login_required
def get_profile_picture(user_id):
    conn = get_db_connection()
    user = conn.execute('SELECT profile_picture FROM users WHERE id = ?', (user_id,)).fetchone()
    conn.close()

    if user and user['profile_picture']:
        return send_file(io.BytesIO(user['profile_picture']), mimetype='image/jpeg')
    else:
        return redirect('/static/imagem/icone.svg')























# #Para Testar






# def add_collaborator_to_idea(idea_id, user_id):
#     try:
#         conn = sqlite3.connect('database.db')
#         c = conn.cursor()
#         c.execute('INSERT INTO idea_users (idea_id, user_id) VALUES (?, ?)', (idea_id, user_id))
#         conn.commit()
#         print("Colaborador adicionado com sucesso!")
#     except sqlite3.IntegrityError:
#         print("Este colaborador já está associado a esta ideia.")
#     except sqlite3.Error as e:
#         print(f"Erro ao adicionar colaborador ao banco de dados: {e}")
#     finally:
#         conn.close()


# def remove_collaborator_from_idea(idea_id, user_id):
#     try:
#         conn = sqlite3.connect('database.db')
#         c = conn.cursor()
#         c.execute('DELETE FROM idea_users WHERE idea_id=? AND user_id=?', (idea_id, user_id))
#         conn.commit()
#         print("Colaborador removido com sucesso!")
#     except sqlite3.Error as e:
#         print(f"Erro ao remover colaborador do banco de dados: {e}")
#     finally:
#         conn.close()



# @app.route('/add_collaborators', methods=['POST'])
# def add_collaborators():
#     idea_id = request.form.get('idea_id')
#     user_ids = request.form.get('user_ids').split(',')
    
#     for user_id in user_ids:
#         user_id = user_id.strip()  # Remove espaços em branco
#         add_collaborator_to_idea(idea_id, user_id)  # Função definida anteriormente
    
#     return redirect(url_for('success'))


# @app.route('/success')
# def success():
#     return "Colaboradores adicionados com sucesso!"


# @app.route('/send_request', methods=['POST'])
# def send_request():
#     idea_id = request.form.get('idea_id')
#     requester_id = request.form.get('requester_id')
#     send_collaboration_request(idea_id, requester_id)  # Função definida anteriormente
#     return redirect(url_for('success'))


# @app.route('/manage_requests', methods=['POST'])
# def manage_requests():
#     idea_id = request.form.get('idea_id')
#     request_id = request.form.get('request_id')
#     action = request.form.get('action')
    
#     if action == 'accept':
#         accept_collaboration_request(request_id)  # Função definida anteriormente
#     elif action == 'reject':
#         reject_collaboration_request(request_id)  # Função definida anteriormente
    
#     return redirect(url_for('success'))


# def reject_collaboration_request(request_id):
#     try:
#         conn = sqlite3.connect('database.db')
#         c = conn.cursor()
#         c.execute('UPDATE collaboration_requests SET status = ? WHERE id = ?', ('recusada', request_id))
#         conn.commit()
#         print("Solicitação de colaboração recusada com sucesso!")
#     except sqlite3.Error as e:
#         print(f"Erro ao recusar solicitação de colaboração: {e}")
#     finally:
#         conn.close()


# def accept_collaboration_request(request_id):
#     try:
#         conn = sqlite3.connect('database.db')
#         c = conn.cursor()
        
#         # Marcar a solicitação como aceita
#         c.execute('UPDATE collaboration_requests SET status = ? WHERE id = ?', ('aceita', request_id))
        
#         # Obter o ID da ideia e o ID do solicitante
#         c.execute('SELECT idea_id, requester_id FROM collaboration_requests WHERE id = ?', (request_id,))
#         idea_id, user_id = c.fetchone()
        
#         # Adicionar o colaborador à ideia
#         add_collaborator_to_idea(idea_id, user_id)
        
#         conn.commit()
#         print("Solicitação de colaboração aceita e colaborador adicionado com sucesso!")
#     except sqlite3.Error as e:
#         print(f"Erro ao aceitar solicitação de colaboração: {e}")
#     finally:
#         conn.close()


# def send_collaboration_request(idea_id, requester_id):
#     try:
#         conn = sqlite3.connect('database.db')
#         c = conn.cursor()
#         c.execute('INSERT INTO collaboration_requests (idea_id, requester_id, status) VALUES (?, ?, ?)', 
#                   (idea_id, requester_id, 'pendente'))
#         conn.commit()
#         print("Solicitação de colaboração enviada com sucesso!")
#     except sqlite3.Error as e:
#         print(f"Erro ao enviar solicitação de colaboração: {e}")
#     finally:
#         conn.close()


if __name__ == '__main__':
    app.run(debug=True)


